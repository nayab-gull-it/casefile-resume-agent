"""
file_generator.py
- extract_text(): pulls raw text out of an uploaded .pdf/.docx/.txt resume
- render_docx() / render_pdf(): turn structured resume JSON into a downloadable file
"""
import os
import pdfplumber
import docx as docx_reader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

INK = "1B2430"
BRASS = "9C7A3C"


def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)
    elif ext == "docx":
        doc = docx_reader.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif ext == "txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def render_docx(resume, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.get("name", "").upper() or "YOUR NAME")
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor.from_string(INK)

    contact = resume.get("contact", {})
    contact_bits = [contact.get("email", ""), contact.get("phone", ""), contact.get("location", "")]
    contact_bits += contact.get("links", []) or []
    contact_line = "  |  ".join([c for c in contact_bits if c])
    if contact_line:
        cp = doc.add_paragraph()
        cr = cp.add_run(contact_line)
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = RGBColor.from_string("555555")

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text.upper())
        r.font.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor.from_string(BRASS)
        p.paragraph_format.border = None
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), BRASS)
        pbdr.append(bottom)
        pPr.append(pbdr)

    if resume.get("summary"):
        add_heading("Summary")
        doc.add_paragraph(resume["summary"])

    if resume.get("skills"):
        add_heading("Skills")
        doc.add_paragraph(" · ".join(resume["skills"]))

    if resume.get("experience"):
        add_heading("Experience")
        for job in resume["experience"]:
            p = doc.add_paragraph()
            r1 = p.add_run(f'{job.get("title","")} — {job.get("company","")}')
            r1.font.bold = True
            r1.font.size = Pt(10.5)
            if job.get("dates"):
                p.add_run(f'   ({job["dates"]})').font.size = Pt(9.5)
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(b, style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)

    if resume.get("projects"):
        add_heading("Projects")
        for proj in resume["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).font.bold = True
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            for b in proj.get("bullets", []):
                doc.add_paragraph(b, style="List Bullet")

    if resume.get("education"):
        add_heading("Education")
        for edu in resume["education"]:
            p = doc.add_paragraph()
            p.add_run(f'{edu.get("degree","")}, {edu.get("institution","")}').font.bold = True
            if edu.get("dates"):
                p.add_run(f'   ({edu["dates"]})').font.size = Pt(9.5)

    doc.save(out_path)
    return out_path


def render_pdf(resume, out_path):
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=20, textColor=HexColor("#" + INK), spaceAfter=2, alignment=0)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9, textColor=HexColor("#555555"), spaceAfter=10)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], fontSize=11.5, textColor=HexColor("#" + BRASS), spaceBefore=12, spaceAfter=4, borderWidth=0)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14)
    job_title_style = ParagraphStyle("JobTitle", parent=styles["Normal"], fontSize=10.5, leading=14, spaceBefore=6)

    doc = SimpleDocTemplate(out_path, pagesize=LETTER,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch,
                             leftMargin=0.7 * inch, rightMargin=0.7 * inch)
    story = []
    story.append(Paragraph(resume.get("name", "Your Name").upper(), name_style))

    contact = resume.get("contact", {})
    contact_bits = [contact.get("email", ""), contact.get("phone", ""), contact.get("location", "")]
    contact_bits += contact.get("links", []) or []
    contact_line = "  |  ".join([c for c in contact_bits if c])
    if contact_line:
        story.append(Paragraph(contact_line, contact_style))

    if resume.get("summary"):
        story.append(Paragraph("SUMMARY", heading_style))
        story.append(Paragraph(resume["summary"], body_style))

    if resume.get("skills"):
        story.append(Paragraph("SKILLS", heading_style))
        story.append(Paragraph(" &middot; ".join(resume["skills"]), body_style))

    if resume.get("experience"):
        story.append(Paragraph("EXPERIENCE", heading_style))
        for job in resume["experience"]:
            title_line = f'<b>{job.get("title","")} — {job.get("company","")}</b>'
            if job.get("dates"):
                title_line += f'  <font size=9 color="#666666">({job["dates"]})</font>'
            story.append(Paragraph(title_line, job_title_style))
            bullets = job.get("bullets", [])
            if bullets:
                story.append(ListFlowable(
                    [ListItem(Paragraph(b, body_style), leftIndent=6) for b in bullets],
                    bulletType="bullet", start="•", leftIndent=14
                ))

    if resume.get("projects"):
        story.append(Paragraph("PROJECTS", heading_style))
        for proj in resume["projects"]:
            story.append(Paragraph(f'<b>{proj.get("name","")}</b>', job_title_style))
            if proj.get("description"):
                story.append(Paragraph(proj["description"], body_style))
            bullets = proj.get("bullets", [])
            if bullets:
                story.append(ListFlowable(
                    [ListItem(Paragraph(b, body_style), leftIndent=6) for b in bullets],
                    bulletType="bullet", start="•", leftIndent=14
                ))

    if resume.get("education"):
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in resume["education"]:
            line = f'<b>{edu.get("degree","")}, {edu.get("institution","")}</b>'
            if edu.get("dates"):
                line += f'  <font size=9 color="#666666">({edu["dates"]})</font>'
            story.append(Paragraph(line, job_title_style))

    doc.build(story)
    return out_path
